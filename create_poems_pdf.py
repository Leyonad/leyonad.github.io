import json
import os
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Konfiguration ---
JSON_PATH = "c:\\Users\\Daniel\\Desktop\\Programmieren\\leyonad.github.io\\stories.json"
BASE_PATH = "c:\\Users\\Daniel\\Desktop\\Programmieren\\leyonad.github.io"
OUTPUT_PATH = os.path.join(BASE_PATH, "gedichte_sammlung_final.docx")
TEMPLATE_PATH = "template.docx" # Muss im selben Ordner wie das Skript liegen!

# --- Hilfsfunktionen ---

def convert_image_to_jpg(image_path):
    if not image_path.lower().endswith((".jpg", ".jpeg")):
        jpg_path = os.path.splitext(image_path)[0] + ".jpg"
        if not os.path.exists(jpg_path):
            try:
                with Image.open(image_path) as img:
                    if img.mode in ("RGBA", "P"):
                        img = img.convert("RGB")
                    img.save(jpg_path, "JPEG")
                print(f"Bild konvertiert: {image_path} -> {jpg_path}")
                return jpg_path
            except Exception as e:
                print(f"Fehler bei der Konvertierung von {image_path}: {e}")
                return None
        else:
            return jpg_path
    return image_path

def add_table_of_contents(doc):
    """
    Fügt einen Platzhalter für ein Inhaltsverzeichnis in das Dokument ein.
    Dieser muss in Word manuell aktualisiert werden.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    
    run = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    # Die Switches bedeuten: \o "1-2" = nimm Überschriften Level 1-2, \h = als Hyperlinks, \z = ohne Seitenzahlen in der Web-Ansicht
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u' 
    run._r.append(instrText)

    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)

    # Platzhalter-Text, bis das Feld aktualisiert wird
    run = paragraph.add_run('Inhaltsverzeichnis')
    run.bold = True
    
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)


# --- Lade die JSON-Datei ---
with open(JSON_PATH, "r", encoding="utf-8") as file:
    data = json.load(file)

# --- Filtere und sortiere die Gedichte ---
poems_filtered = [
    item for item in data["items"] 
    if item["type"] == "poem" and item.get("allow_sending") is True
]
poems_sorted = sorted(poems_filtered, key=lambda p: p["started"], reverse=True)


# --- Erstelle das Word-Dokument ---
try:
    document = Document(TEMPLATE_PATH)
except Exception as e:
    print(f"FEHLER: Die Vorlagendatei '{TEMPLATE_PATH}' konnte nicht gefunden werden.")
    exit()

# --- NEU: Angepasste Struktur für Titelseite und Inhaltsverzeichnis ---
# 1. Haupttitel (verwendet die "Untertitel"-Formatvorlage aus dem Template)
document.add_paragraph("Gedichte von Daniel Holzknecht", style="Untertitel")

# 2. Inhaltsverzeichnis
add_table_of_contents(document)
document.add_page_break()

# --- Füge die Gedichte hinzu ---
for i, poem in enumerate(poems_sorted):
    print(f"Verarbeite Gedicht (Datum: {poem['started']}): {poem['title']}...")

    # Titel (verwendet die "Überschrift 2"-Formatvorlage, erscheint im Inhaltsverzeichnis)
    document.add_heading(poem["title"], level=2)
    
    # # Bild
    # image_path = os.path.join(BASE_PATH, poem["image"])
    # converted_image_path = convert_image_to_jpg(image_path)
    # if converted_image_path and os.path.exists(converted_image_path):
    #     document.add_picture(converted_image_path, width=Inches(2.0))
    # else:
    #     document.add_paragraph("[Bild nicht gefunden]").italic = True
    
    # document.add_paragraph()

    # Gedichttext
    content_path = os.path.join(BASE_PATH, poem["contentPath"])
    if os.path.exists(content_path):
        with open(content_path, "r", encoding="utf-8") as content_file:
            lines = content_file.readlines()
            for line in lines:
                if line.strip().startswith('#'): continue
                document.add_paragraph(line.rstrip('\n'))
    else:
        document.add_paragraph("[Gedichttext nicht gefunden]").italic = True

    if i < len(poems_sorted) - 1:
        document.add_page_break()

# --- Speichere das finale Dokument ---
try:
    document.save(OUTPUT_PATH)
    print("-" * 30)
    print(f"Word-Dokument wurde erfolgreich erstellt: {OUTPUT_PATH}")
    print(f"{len(poems_sorted)} Gedichte wurden verarbeitet.")
    print("\n--- WICHTIGER HINWEIS ---")
    print("Öffne die Datei in Word und aktualisiere das Inhaltsverzeichnis:")
    print("1. Rechtsklick auf den Text 'Inhaltsverzeichnis'.")
    print("2. Wähle 'Felder aktualisieren'.")
    print("3. Wähle 'Gesamtes Verzeichnis aktualisieren' und klicke OK.")

except Exception as e:
    print(f"FEHLER beim Speichern des Dokuments: {e}")